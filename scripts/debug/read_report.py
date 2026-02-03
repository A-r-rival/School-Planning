
import docx
import sys
import os

file_path = os.path.join("A. Donem 1 Rapor", "OTOMATİK ÜNIVERSITE DERS PROGRAMI OLUŞTURMA SİSTEMİ.docx")

try:
    doc = docx.Document(file_path)
    
    with open("report_full.txt", "w", encoding="utf-8") as f:
        f.write(f"Reading file: {file_path}\n\n")
        f.write("-" * 50 + "\n")
        
        for para in doc.paragraphs:
            if para.text.strip():
                f.write(para.text + "\n\n")
                
        f.write("-" * 50 + "\n")
        f.write("TABLES CONTENT:\n")
        for table in doc.tables:
            for row in table.rows:
                row_text = [cell.text.strip() for cell in row.cells]
                f.write(" | ".join(row_text) + "\n")
            f.write("-" * 20 + "\n")
            
    print("Successfully wrote to report_full.txt")
        
except Exception as e:
    print(f"Error reading docx: {e}")
